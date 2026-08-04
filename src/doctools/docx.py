"""Word 文档（.docx）处理逻辑。"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def _clear_defined_part(part) -> None:
    """清空一个已定义（未继承）页眉/页脚的所有内容。

    包括文字、图片等 run 元素，以及表格。被其他节继承的部分无需单独
    处理——它们继承自本部分，来源清空后自然随之清空。段落直接格式里的
    边框（``w:pBdr``，页眉下方/页脚上方的横线）属于段落格式而非文字，
    单独清空 run 不会移除，需一并删掉。
    """
    for paragraph in list(part.paragraphs):
        for run in list(paragraph.runs):
            run._r.getparent().remove(run._r)

        pPr = paragraph._p.find(qn("w:pPr"))
        if pPr is not None:
            pBdr = pPr.find(qn("w:pBdr"))
            if pBdr is not None:
                pPr.remove(pBdr)

    for table in list(part.tables):
        table._tbl.getparent().remove(table._tbl)


def _remove_style_border(doc: Document, style_ids: set[str], names: set[str]) -> None:
    """从匹配的段落样式定义里移除段落边框。

    Word 页眉/页脚文字旁的横线常来自 styles.xml 中 Header / Footer 样式的
    ``w:pBdr``（页眉取下边框 w:bottom、页脚取上边框 w:top），而非段落直接
    格式。即使清空了内容，只要段落仍引用该样式，横线就会保留，因此需一并
    从样式定义中移除。
    """
    styles_el = doc.styles.element
    for style in styles_el.findall(qn("w:style")):
        style_id = style.get(qn("w:styleId"))
        name_el = style.find(qn("w:name"))
        name = name_el.get(qn("w:val")) if name_el is not None else None
        if style_id not in style_ids and name not in names:
            continue
        pPr = style.find(qn("w:pPr"))
        if pPr is None:
            continue
        pBdr = pPr.find(qn("w:pBdr"))
        if pBdr is not None:
            pPr.remove(pBdr)


def _remove_header_style_border(doc: Document) -> None:
    """移除“Header”段落样式定义里的段落边框。"""
    _remove_style_border(doc, {"Header"}, {"header"})


def _remove_footer_style_border(doc: Document) -> None:
    """移除“Footer”段落样式定义里的段落边框。"""
    _remove_style_border(doc, {"Footer"}, {"footer"})


def clear_headers(doc: Document) -> None:
    """去除文档所有节的页眉（含首页、奇偶页变体）。

    只处理那些自带内容的页眉（``is_linked_to_previous`` 为 False 的），
    避免为不使用的首页 / 奇偶页页眉生成多余的 XML 定义。
    """
    _remove_header_style_border(doc)
    for section in doc.sections:
        for header in (
            section.header,
            section.first_page_header,
            section.even_page_header,
        ):
            if not header.is_linked_to_previous:
                _clear_defined_part(header)


def strip_headers(src: Path, dst: Path) -> None:
    """读取一个 .docx，去除页眉后保存到 dst。"""
    doc = Document(str(src))
    clear_headers(doc)
    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst))


def clear_footers(doc: Document) -> None:
    """去除文档所有节的页脚（含首页、奇偶页变体）。"""
    _remove_footer_style_border(doc)
    for section in doc.sections:
        for footer in (
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            if not footer.is_linked_to_previous:
                _clear_defined_part(footer)


def strip_footers(src: Path, dst: Path) -> None:
    """读取一个 .docx，去除页脚后保存到 dst。"""
    doc = Document(str(src))
    clear_footers(doc)
    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst))


def clear_headers_footers(doc: Document) -> None:
    """同时去除文档所有节的页眉与页脚。"""
    clear_headers(doc)
    clear_footers(doc)


def strip_headers_footers(src: Path, dst: Path) -> None:
    """读取一个 .docx，同时去除页眉页脚后保存到 dst。"""
    doc = Document(str(src))
    clear_headers_footers(doc)
    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst))
