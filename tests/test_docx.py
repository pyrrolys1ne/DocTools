from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from typer.testing import CliRunner

from doctools.batch import build_plan
from doctools.cli import app
from doctools.docx import clear_headers, strip_headers

runner = CliRunner()


def _doc_with_headers() -> Document:
    doc = Document()

    # 第一节：普通页眉（多行文字）
    h1 = doc.sections[0].header
    h1.is_linked_to_previous = False
    h1.paragraphs[0].text = "CONFIDENTIAL"
    h1.add_paragraph("第二行页眉")
    doc.add_paragraph("正文第一段")

    # 第二节：自定义页眉
    s2 = doc.add_section()
    h2 = s2.header
    h2.is_linked_to_previous = False
    h2.paragraphs[0].text = "Section2 header"
    doc.add_paragraph("正文第二段")
    return doc


def test_clear_headers_removes_all_header_text() -> None:
    doc = _doc_with_headers()
    clear_headers(doc)
    for section in doc.sections:
        assert all(p.text == "" for p in section.header.paragraphs)
        assert all(p.text == "" for p in section.first_page_header.paragraphs)
        assert all(p.text == "" for p in section.even_page_header.paragraphs)


def test_clear_headers_keeps_body_text() -> None:
    doc = _doc_with_headers()
    clear_headers(doc)
    body = [p.text for p in doc.paragraphs]
    assert "正文第一段" in body
    assert "正文第二段" in body


def test_clear_headers_handles_first_page_header() -> None:
    doc = _doc_with_headers()
    doc.sections[0].different_first_page_header_footer = True
    first = doc.sections[0].first_page_header
    first.is_linked_to_previous = False
    first.paragraphs[0].text = "FIRST PAGE"
    clear_headers(doc)
    assert doc.sections[0].first_page_header.paragraphs[0].text == ""


def test_clear_headers_removes_paragraph_border() -> None:
    """页眉段落的下边框（w:pBdr）是段落格式，去页眉时应一并移除。"""
    doc = _doc_with_headers()
    header = doc.sections[0].header
    header.is_linked_to_previous = False
    pPr = header.paragraphs[0]._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pPr.makeelement(qn("w:bottom"), {qn("w:val"): "single", qn("w:sz"): "4"})
    pBdr.append(bottom)
    pPr.append(pBdr)

    clear_headers(doc)

    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            for paragraph in header.paragraphs:
                _pPr = paragraph._p.find(qn("w:pPr"))
                assert _pPr is None or _pPr.find(qn("w:pBdr")) is None


def _header_style_with_border(doc: Document) -> None:
    """给文档的 Header 段落样式注入下边框（Word 页眉横线的来源）。"""
    for style in doc.styles.element.findall(qn("w:style")):
        if style.get(qn("w:styleId")) != "Header":
            continue
        pPr = style.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pPr.makeelement(qn("w:bottom"), {qn("w:val"): "single", qn("w:sz"): "6"})
        pBdr.append(bottom)
        pPr.insert(0, pBdr)
        return


def test_clear_headers_removes_header_style_border() -> None:
    """Word 页眉横线常来自 Header 样式的 w:pBdr，需从样式定义中一并移除。"""
    doc = _doc_with_headers()
    _header_style_with_border(doc)

    clear_headers(doc)

    for style in doc.styles.element.findall(qn("w:style")):
        if style.get(qn("w:styleId")) != "Header":
            continue
        pPr = style.find(qn("w:pPr"))
        assert pPr is None or pPr.find(qn("w:pBdr")) is None


def test_strip_headers_file(tmp_path: Path) -> None:
    src = tmp_path / "in.docx"
    dst = tmp_path / "out.docx"
    _doc_with_headers().save(str(src))

    strip_headers(src, dst)

    out = Document(str(dst))
    assert all(p.text == "" for s in out.sections for p in s.header.paragraphs)
    assert any(p.text.startswith("正文") for p in out.paragraphs)


def test_cli_remove_headers_single_file(tmp_path: Path) -> None:
    src = tmp_path / "in.docx"
    dst = tmp_path / "out.docx"
    _doc_with_headers().save(str(src))

    result = runner.invoke(app, ["remove-headers", str(src), "--output", str(dst)])

    assert result.exit_code == 0, result.stdout
    out = Document(str(dst))
    assert all(p.text == "" for s in out.sections for p in s.header.paragraphs)


def test_cli_remove_headers_dry_run(tmp_path: Path) -> None:
    src = tmp_path / "in.docx"
    _doc_with_headers().save(str(src))

    result = runner.invoke(app, ["remove-headers", str(src), "--dry-run"])

    assert result.exit_code == 0
    assert "[dry-run]" in result.stdout
    assert not (tmp_path / "in_cleaned.docx").exists()


def test_cli_rejects_non_docx(tmp_path: Path) -> None:
    src = tmp_path / "notes.txt"
    src.write_text("hi")

    result = runner.invoke(app, ["remove-headers", str(src)])

    assert result.exit_code != 0
    assert "仅支持 .docx" in result.stderr


def test_build_plan_single_file_output_is_full_path(tmp_path: Path) -> None:
    src = tmp_path / "in.docx"
    src.write_bytes(b"x")
    out = tmp_path / "renamed.docx"

    pairs = build_plan(src, out)

    assert pairs == [(src, out)]


def test_build_plan_single_file_output_is_dir(tmp_path: Path) -> None:
    src = tmp_path / "in.docx"
    src.write_bytes(b"x")
    out_dir = tmp_path / "out"

    pairs = build_plan(src, out_dir, output_is_dir=True)

    assert pairs == [(src, out_dir / "in_cleaned.docx")]


def test_build_plan_recursive_mirrors_structure(tmp_path: Path) -> None:
    src = tmp_path / "in"
    (src / "sub").mkdir(parents=True)
    _doc_with_headers().save(str(src / "a.docx"))
    _doc_with_headers().save(str(src / "sub" / "b.docx"))

    pairs = build_plan(src, tmp_path / "out", recursive=True)

    assert pairs == [
        (src / "a.docx", tmp_path / "out" / "a.docx"),
        (src / "sub" / "b.docx", tmp_path / "out" / "sub" / "b.docx"),
    ]


def test_cli_remove_headers_recursive(tmp_path: Path) -> None:
    src = tmp_path / "in"
    (src / "sub").mkdir(parents=True)
    _doc_with_headers().save(str(src / "a.docx"))
    _doc_with_headers().save(str(src / "sub" / "b.docx"))

    result = runner.invoke(
        app, ["remove-headers", str(src), "-o", str(tmp_path / "out"), "-r"]
    )

    assert result.exit_code == 0, result.stdout
    for out in (tmp_path / "out" / "a.docx", tmp_path / "out" / "sub" / "b.docx"):
        assert out.exists()
        doc = Document(str(out))
        assert all(p.text == "" for s in doc.sections for p in s.header.paragraphs)


def test_cli_remove_headers_non_recursive_skips_subdirs(tmp_path: Path) -> None:
    src = tmp_path / "in"
    (src / "sub").mkdir(parents=True)
    _doc_with_headers().save(str(src / "a.docx"))
    _doc_with_headers().save(str(src / "sub" / "b.docx"))

    result = runner.invoke(app, ["remove-headers", str(src), "-o", str(tmp_path / "out")])

    assert result.exit_code == 0, result.stdout
    assert (tmp_path / "out" / "a.docx").exists()
    assert not (tmp_path / "out" / "sub" / "b.docx").exists()
