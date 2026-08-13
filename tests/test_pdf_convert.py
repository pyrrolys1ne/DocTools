"""PDF → Word / PDF → PPT 转换测试。"""

from __future__ import annotations

import zipfile
from pathlib import Path

import pytest
from docx import Document
from pptx import Presentation
from reportlab.pdfgen import canvas

from doctools.batch import build_convert_plan
from doctools.errors import PDF_NO_TEXT, DoctoolsError
from doctools.pdf_convert import pdf_to_docx, pdf_to_images, pdf_to_pptx


def _make_pdf(path: Path, pages: int = 2) -> None:
    c = canvas.Canvas(str(path))
    for i in range(1, pages + 1):
        c.drawString(72, 720, f"Test page {i}")
        c.showPage()
    c.save()


def test_pdf_to_docx_contains_text(tmp_path: Path) -> None:
    src = tmp_path / "in.pdf"
    dst = tmp_path / "out.docx"
    _make_pdf(src)

    pdf_to_docx(src, dst)

    assert dst.exists()
    text = " ".join(p.text for p in Document(str(dst)).paragraphs)
    assert "Test" in text


def test_pdf_to_pptx_one_slide_per_page(tmp_path: Path) -> None:
    src = tmp_path / "in.pdf"
    dst = tmp_path / "out.pptx"
    _make_pdf(src, pages=3)

    pdf_to_pptx(src, dst)

    assert dst.exists()
    assert len(Presentation(str(dst)).slides) == 3


def test_pdf_to_images_one_png_per_page(tmp_path: Path) -> None:
    src = tmp_path / "in.pdf"
    out = tmp_path / "out"
    _make_pdf(src, pages=2)

    results = pdf_to_images(src, out)

    assert all(r.ok for r in results)
    assert sorted(p.name for p in out.glob("*.png")) == ["in_p1.png", "in_p2.png"]


def test_build_convert_plan_pdf_to_word_naming(tmp_path: Path) -> None:
    src = tmp_path / "in.pdf"
    _make_pdf(src)

    pairs = build_convert_plan(
        src, tmp_path / "out", output_is_dir=True,
        suffixes=(".pdf",), out_suffix=".docx", default_out_dir="docx",
    )

    assert pairs == [(src, tmp_path / "out" / "in.docx")]


def test_pdf_to_docx_fallback_on_engine_failure(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """pdf2docx 引擎失败时回退到文字提取，输出仍是合法 docx 并返回附注。"""
    src = tmp_path / "in.pdf"
    dst = tmp_path / "out.docx"
    _make_pdf(src)

    def _boom(*_args, **_kwargs) -> None:
        raise RuntimeError("pdf2docx exploded")

    monkeypatch.setattr("doctools.pdf_convert._convert_with_pdf2docx", _boom)

    note = pdf_to_docx(src, dst)

    assert note is not None and "回退" in note
    assert dst.exists()
    with zipfile.ZipFile(str(dst)) as zf:  # docx 是合法 OOXML zip
        assert "word/document.xml" in zf.namelist()
    assert any("Test" in p.text for p in Document(str(dst)).paragraphs)


def test_pdf_to_docx_scanned_pdf_reports_no_text(tmp_path: Path) -> None:
    """纯图片（无文字）PDF 报 PDF_NO_TEXT，而不是产出空白 docx。"""
    src = tmp_path / "scan.pdf"
    c = canvas.Canvas(str(src))
    c.rect(50, 50, 500, 700, fill=1, stroke=0)  # 只画色块，无文字
    c.showPage()
    c.save()

    with pytest.raises(DoctoolsError) as excinfo:
        pdf_to_docx(src, tmp_path / "out.docx")

    assert excinfo.value.code == PDF_NO_TEXT
